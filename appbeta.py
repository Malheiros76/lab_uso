import streamlit as st
from pymongo import MongoClient
import pandas as pd
from datetime import datetime
import pytz  # import pytz para fuso horário
import io
import base64
from fpdf import FPDF
from docx import Document
import qrcode
from PIL import Image
from docx.shared import Inches


# --- Fuso horário São Paulo ---
tz_sao_paulo = pytz.timezone('America/Buenos_Aires')

# --- MongoDB connection ---
client = MongoClient("mongodb+srv://bibliotecaluizcarlos:FGVIlVhDcUDuQddG@cluster0.hth6xs5.mongodb.net/?retryWrites=true&w=majority")
db = client["controle_uso"]

# --- Funções auxiliares ---
def generate_qrcode(data):
    qr = qrcode.QRCode(box_size=5, border=2)
    qr.add_data(data)
    qr.make(fit=True)
    img = qr.make_image(fill_color="black", back_color="white")
    return img

def pil_image_to_bytes(img):
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()

def importar_alunos_txt(arquivo):
    try:
        df = pd.read_csv(arquivo, sep="\t", dtype=str)
        for _, row in df.iterrows():
            aluno_doc = {
                "cgm": str(row.get("CGM", "")).strip(),
                "nome": str(row.get("Nome do Estudante", "")).strip(),
                "turma": str(row.get("Turma", "")).strip(),
                "data_nasc": str(row.get("Data de Nasc.", "")).strip(),
                "sexo": str(row.get("Sexo", "")).strip(),
                "telefone": str(row.get("Telefone", "")).strip()
            }
            db.alunos.update_one({"cgm": aluno_doc["cgm"]}, {"$set": aluno_doc}, upsert=True)
        st.success("Alunos importados com sucesso!")
    except Exception as e:
        st.error(f"Erro ao processar o arquivo: {e}")

# --- Login ---
def login():
    st.sidebar.title("🔐 Controle de Uso")
    usuario = st.sidebar.text_input("👤 Usuário", key="login_usuario")
    senha = st.sidebar.text_input("🔑 Senha", type="password", key="login_senha")
    if st.sidebar.button("✅ Login"):
        user = db.usuarios.find_one({"usuario": usuario, "senha": senha})
        if user:
            st.session_state["logado"] = True
            st.session_state["usuario"] = usuario
            st.session_state["usuario_admin"] = user.get("nivel", "user") == "admin"
            st.rerun()
        else:
            st.sidebar.error("Usuário ou senha incorretos")

def cadastro_admin():
    st.sidebar.title("🔧 Primeiro Acesso - Crie o Admin")
    usuario = st.text_input("👤 Usuário admin", key="cadastro_admin_usuario")
    senha = st.text_input("🔑 Senha admin", type="password", key="cadastro_admin_senha")
    if st.button("✅ Cadastrar Usuário Admin"):
        if usuario and senha:
            existe = db.usuarios.find_one({"usuario": usuario})
            if existe:
                st.warning("Usuário já existe!")
            else:
                db.usuarios.insert_one({"usuario": usuario, "senha": senha, "nivel": "admin"})
                st.success("Usuário admin criado com sucesso! Faça login.")
                st.stop()
        else:
            st.warning("Preencha usuário e senha.")

# --- Cadastro Gerais ---
def cadastro_gerais():
    aba = st.selectbox("🗂️ Selecione o cadastro", [
        "Usuários", "Laboratórios", "Mesas", "Equipamentos", "Importação de Alunos"
    ])

    if aba == "Usuários":
        st.subheader("👥 Cadastro de Usuários")
        usuario = st.text_input("Usuário", key="cadastro_usuario")
        senha = st.text_input("Senha", type="password", key="cadastro_senha")
        nivel = st.selectbox("Nível", ["admin", "user"], key="cadastro_nivel")
        if st.button("💾 Salvar Usuário"):
            if usuario and senha:
                existe = db.usuarios.find_one({"usuario": usuario})
                if existe:
                    st.warning("Usuário já existe!")
                else:
                    db.usuarios.insert_one({"usuario": usuario, "senha": senha, "nivel": nivel})
                    st.success("Usuário salvo!")
            else:
                st.warning("Preencha todos os campos.")

    elif aba == "Laboratórios":
        st.subheader("🔬 Cadastro de Laboratórios")
        nome = st.text_input("Nome do laboratório", key="cadastro_lab")
        if st.button("💾 Salvar Laboratório"):
            if nome:
                db.laboratorios.insert_one({"nome": nome})
                st.success("Laboratório salvo!")
            else:
                st.warning("Preencha o nome.")

    elif aba == "Mesas":
        st.subheader("🪑 Cadastro de Mesas")
        if st.button("🔄 Cadastrar mesas de 1 a 40"):
            db.mesas.delete_many({})
            for i in range(1, 41):
                db.mesas.insert_one({"numero": str(i)})
            st.success("Mesas cadastradas com sucesso!")

        numero = st.text_input("Número da Mesa", key="numero_mesa")
        if st.button("💾 Salvar Mesa"):
            if numero:
                db.mesas.insert_one({"numero": numero})
                st.success("Mesa salva!")
            else:
                st.warning("Preencha o número.")

    elif aba == "Equipamentos":
        st.subheader("💻 Cadastro de Equipamentos")
        if st.button("🔄 Cadastrar equipamentos padrão"):
            db.equipamentos.delete_many({})
            equipamentos = {"Tablet": 40, "Chromebook": 35, "CPU": 24, "Netbook": 27}
            for nome, quantidade in equipamentos.items():
                for i in range(1, quantidade + 1):
                    db.equipamentos.insert_one({"nome": nome, "numero": str(i)})
            st.success("Equipamentos cadastrados!")

        nome = st.text_input("Nome do equipamento")
        numero = st.text_input("Número do equipamento")
        if st.button("💾 Salvar Equipamento"):
            if nome and numero:
                db.equipamentos.insert_one({"nome": nome, "numero": numero})
                st.success("Equipamento salvo!")
            else:
                st.warning("Preencha todos os campos.")

    elif aba == "Importação de Alunos":
        st.subheader("📥 Importação de Alunos")
        arquivo = st.file_uploader("Importar arquivo .txt", type="txt")
        if arquivo:
            importar_alunos_txt(arquivo)

def imprimir_qrcodes():
    st.subheader("🖨️ Imprimir QR Codes")

    opcao = st.radio("Escolha o que deseja imprimir:", ["Mesas", "Equipamentos"])

    if opcao == "Mesas":
        dados = list(db.mesas.find())
        titulo = "Mesa"
    else:
        dados = list(db.equipamentos.find())
        titulo = "Equipamento"

    if not dados:
        st.info(f"Nenhum dado cadastrado para {titulo.lower()}s.")
        return

    imagens = []
    faltando_info = False
    for item in dados:
        if titulo == "Mesa":
            numero = item.get("numero", None)
            texto = f"Mesa {numero}" if numero else "Mesa ?"
        else:
            nome = item.get("nome", None)
            numero = item.get("numero", None)
            texto = f"{nome} #{numero}" if nome and numero else f"{nome if nome else 'Equipamento'} #{numero if numero else '?'}"
        img = generate_qrcode(texto)
        imagens.append((texto, img))

    if faltando_info:
        st.warning("Alguns itens estão com informações incompletas e foram mostrados com '?'.")

    st.write("🔲 QR Codes Gerados:")
    cols = st.columns(3)
    for i, (texto, img) in enumerate(imagens):
        with cols[i % 3]:
            st.image(pil_image_to_bytes(img), caption=texto, width=120)

    # PDF A4
    pdf = FPDF(orientation="P", unit="mm", format="A4")
    pdf.set_auto_page_break(auto=True, margin=10)
    pdf.add_page()
    pdf.set_font("Arial", size=10)

    num_cols = 3
    cell_w = 60
    cell_h = 60
    margin_x = 10
    margin_y = 20
    x = margin_x
    y = margin_y

    for i, (texto, img) in enumerate(imagens):
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        buf.seek(0)

        pdf.image(buf, x=x + 5, y=y, w=cell_w - 10)
        pdf.set_xy(x, y + cell_h - 5)
        pdf.multi_cell(cell_w, 5, texto, align="C")

        x += cell_w
        if (i + 1) % num_cols == 0:
            x = margin_x
            y += cell_h + 10
            if y + cell_h > 280:
                pdf.add_page()
                y = margin_y

    pdf_bytes = bytes(pdf.output(dest="S"))
    b64 = base64.b64encode(pdf_bytes).decode()
    href = f'<a href="data:application/pdf;base64,{b64}" download="qrcodes_{titulo.lower()}.pdf">📥 Baixar QR Codes em PDF (A4)</a>'
    st.markdown(href, unsafe_allow_html=True)

# --- Registro de Uso (selectbox ao invés de QR Code na webcam) ---
def registro_uso():
    st.subheader("📝 Registro de Uso")

    alunos = list(db.alunos.find())
    if not alunos:
        st.warning("Nenhum aluno cadastrado!")
        return

    opcoes = [f"{a['nome']} ({a['turma']})" for a in alunos]
    indice = st.selectbox("Aluno", range(len(opcoes)), format_func=lambda x: opcoes[x])
    aluno_selecionado = alunos[indice]

    # Default time já no fuso São Paulo
    horario = st.time_input("Horário", value=datetime.now(tz_sao_paulo).time())

    mesas = list(db.mesas.find())
    opcoes_mesas = [m["numero"] for m in mesas] if mesas else []
    mesa = st.selectbox("Mesa", [""] + opcoes_mesas)

    equipamentos = list(db.equipamentos.find())
    opcoes_equipamentos = [f"{e['nome']} #{e['numero']}" for e in equipamentos] if equipamentos else []
    equipamento = st.selectbox("Equipamento", [""] + opcoes_equipamentos)

    if st.button("💾 Registrar Uso"):
        if aluno_selecionado and horario and (mesa.strip() or equipamento.strip()):
            # Pega a data/hora atual em UTC para salvar (padrão recomendado)
            now_utc = datetime.utcnow().replace(tzinfo=pytz.UTC)

            db.registros.insert_one({
                "data": now_utc,
                "aluno_cgm": aluno_selecionado["cgm"],
                "aluno_nome": aluno_selecionado["nome"],
                "horario": horario.strftime("%H:%M"),
                "mesa": mesa.strip(),
                "equipamento": equipamento.strip()
            })
            st.success("Uso registrado com sucesso!")
        else:
            st.warning("Preencha todos os campos obrigatórios (aluno, horário e mesa ou equipamento).")

# --- Relatórios ---
def relatorios():
    st.subheader("📊 Relatórios de Uso")

    registros = list(db.registros.find())
    if registros:
        df = pd.DataFrame(registros)

        # Função para converter datas para fuso de São Paulo
        def converter_data(x):
            if isinstance(x, datetime):
                if x.tzinfo is None:
                    x = x.replace(tzinfo=pytz.UTC)
                return x.astimezone(tz_sao_paulo).strftime("%d/%m/%Y %H:%M")
            return x

        if "data" in df.columns:
            df["data"] = df["data"].apply(converter_data)

        st.dataframe(df[["data", "aluno_nome", "horario", "mesa", "equipamento"]])

        col1, col2 = st.columns(2)

        # --- Exportação DOCX ---
        with col1:
            if st.button("📄 Exportar DOCX"):
                doc = Document()
                doc.add_heading('Relatório de Uso', 0)

                # Inserir brasão no topo do DOCX
                try:
                    doc.add_picture("BRASÃO.png", width=Inches(1.5))
                except Exception as e:
                    st.warning(f"⚠️ Não foi possível adicionar o brasão ao DOCX: {e}")

                doc.add_paragraph("")  # espaço

                for _, row in df.iterrows():
                    doc.add_paragraph(
                        f"{row['data']} - {row['aluno_nome']} - {row['horario']} - Mesa {row['mesa']} - {row['equipamento']}"
                    )

                buffer = io.BytesIO()
                doc.save(buffer)
                b64 = base64.b64encode(buffer.getvalue()).decode()
                href = f'<a href="data:application/octet-stream;base64,{b64}" download="relatorio.docx">📥 Baixar DOCX</a>'
                st.markdown(href, unsafe_allow_html=True)

        # --- Exportação PDF ---
       with col2:
    if st.button("🧾 Exportar PDF"):
        pdf = FPDF()
        pdf.add_page()

        try:
            pdf.image("BRASÃO.png", x=10, y=8, w=30)
        except Exception as e:
            st.warning(f"⚠️ Não foi possível adicionar o brasão ao PDF: {e}")

        pdf.set_font("Arial", size=12)
        pdf.cell(200, 10, txt="Relatório de Uso", ln=True, align='C')
        pdf.ln(20)

        for _, row in df.iterrows():
            texto = f"{row['data']} - {row['aluno_nome']} - {row['horario']} - Mesa {row['mesa']} - {row['equipamento']}"
            pdf.cell(0, 10, txt=texto, ln=True)

        # ✅ CORREÇÃO AQUI:
        pdf_bytes = pdf.output(dest="S").encode('latin1')
        b64 = base64.b64encode(pdf_bytes).decode()
        href = f'<a href="data:application/pdf;base64,{b64}" download="relatorio.pdf">📥 Baixar PDF</a>'
        st.markdown(href, unsafe_allow_html=True)
        
    else:
        st.info("Nenhum registro encontrado.")

# --- Main ---
def main():
    st.set_page_config(page_title="Controle de Uso", layout="wide")
    st.sidebar.image("BRASÃO.png", width=120)

    if "logado" not in st.session_state:
        st.session_state["logado"] = False
        st.session_state["usuario_admin"] = False

    if db.usuarios.count_documents({}) == 0 and not st.session_state["logado"]:
        cadastro_admin()
    elif not st.session_state["logado"]:
        login()
    else:
        menu_opcoes = ["📝 Registro de Uso", "📊 Relatórios"]
        if st.session_state.get("usuario_admin"):
            menu_opcoes += ["🗂️ Cadastros Gerais", "🖨️ Impressão QR Codes"]
        menu_opcoes.append("🚪 Sair")

        menu = st.sidebar.radio("Menu", menu_opcoes)

        if menu == "📝 Registro de Uso":
            registro_uso()
        elif menu == "📊 Relatórios":
            relatorios()
        elif menu == "🗂️ Cadastros Gerais":
            cadastro_gerais()
        elif menu == "🖨️ Impressão QR Codes":
            imprimir_qrcodes()
        elif menu == "🚪 Sair":
            if st.sidebar.button("🚪 Clique aqui para Sair"):
                for key in ["logado", "usuario_admin", "usuario", "codigo_lido"]:
                    if key in st.session_state:
                        del st.session_state[key]
                st.rerun()

if __name__ == "__main__":
    main()
